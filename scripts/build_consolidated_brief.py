from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/discovery/Jeff Outreach Engine Consolidated Discovery Brief.docx")
BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 99, 110)
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = "FFFFFF"


def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_cell_text(cell, bold=False, color=None, size=9.5):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08
        for run in paragraph.runs:
            set_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade(cell, LIGHT)
        style_cell_text(cell, bold=True, color=DARK)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            style_cell_text(cells[index])
        set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.left_indent = Inches(0.375 if level == 0 else 0.625)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    set_font(paragraph.add_run(text))
    return paragraph


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, PALE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run(f"{label}: "), bold=True, color=DARK)
    set_font(paragraph.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after, color in (
    ("Heading 1", 16, 18, 8, BLUE),
    ("Heading 2", 13, 14, 7, BLUE),
    ("Heading 3", 11.5, 10, 5, DARK),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(header.add_run("AI GROWTH PLAN  |  JEFF OUTREACH ENGINE"), size=8.5, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("Working discovery brief  •  Updated August 16, 2026"), size=8.5, color=MUTED)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(44)
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("JEFF OUTREACH ENGINE"), size=27, bold=True, color=DARK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Consolidated Discovery Brief & Follow-Up Submission Guide"), size=15, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
set_font(p.add_run("Prepared for Jeff and Coach Bouf  |  Working draft—rules require validation"), size=10.5, color=MUTED, italic=True)
add_note(doc, "Purpose", "Combine the existing discovery materials into one source of truth, distinguish confirmed project direction from proposed starting points, and make Jeff’s next submission easy to complete.")

heading(doc, "1. Executive Summary")
doc.add_paragraph("The existing material establishes a strong discovery and design framework, but it does not yet contain enough completed evidence to build autonomous outreach agents safely. The next phase is to collect Jeff’s real examples, observe his live workflow, and convert his judgment into explicit rules.")
add_bullet(doc, "HubSpot is intended to remain the CRM and operational source of truth.")
add_bullet(doc, "Codex is intended to serve as the AI operating and orchestration layer.")
add_bullet(doc, "The system should support research, qualification, messaging, CRM upkeep, meeting preparation, follow-up, and pipeline coaching.")
add_bullet(doc, "Human review is required before publishing, sending, changing sensitive CRM data, scheduling, or communicating with existing clients until Jeff explicitly approves a narrower autonomy policy.")
add_bullet(doc, "The framework should be reusable for other representatives, while Jeff’s ICP, voice, scoring, and decision rules remain configurable.")
add_note(doc, "Current status", "Discovery framework complete. Jeff-specific sales rules, evidence, examples, system access details, and approval boundaries remain open.")

heading(doc, "2. Source Review and Consolidation")
add_table(doc, ["Source", "What It Contributes", "Disposition"], [
    ("Jeff AI Revenue Engine Sales Process Discovery Interview Questionnaire", "Comprehensive 22-section master question bank covering the complete sales lifecycle.", "Retain as the master discovery reference."),
    ("Outreach Engine Sales Process Discovery Interview Questionnaire", "Tighter version of the same question bank with revised Outreach Engine naming.", "Absorbed; do not maintain separately unless a short interviewer version is needed."),
    ("Outreach Engine Total Llan", "120-minute agenda, proposed funnel, nine-agent model, HubSpot fields, live tests, and assignments.", "Retain as the session/build-plan source; proposals still require validation."),
    ("Jeff P Interview 1 audio", "Approximately 60 minutes of primary-source conversation supplied with this project.", "Preserve as evidence; add a verified transcript and extracted decisions before treating statements as final rules."),
], [2400, 4560, 2400])

heading(doc, "3. Project Mission and Phase 1 Boundary")
doc.add_paragraph("Mission: Build a practical, approval-gated Outreach Engine that helps Jeff identify better prospects, prepare stronger outreach, follow up consistently, and keep HubSpot current—without replacing Jeff’s judgment or creating compliance and trust risks.")
heading(doc, "In Phase 1", 2)
for item in [
    "Capture Jeff’s real end-to-end prospecting process and decision logic.",
    "Define the ICP, qualification rules, opportunity score, decision-maker hierarchy, and research standards.",
    "Draft outreach and follow-up content for Jeff’s review.",
    "Define HubSpot data requirements, attribution fields, and approval-gated updates.",
    "Test the workflow against three to five real prospects.",
    "Measure time saved, accuracy, message quality, and follow-up consistency.",
]: add_bullet(doc, item)
heading(doc, "Outside Phase 1 Until Approved", 2)
for item in [
    "Unreviewed outbound sending or automatic LinkedIn messaging.",
    "Automatic changes to deal stages, ownership, attribution, or commission qualification.",
    "Automatic scheduling or communication with existing clients.",
    "Claims that a prospect is overspending or guaranteed to save money without evidence.",
    "Complex automation architecture before the manual and AI-assisted workflow is validated.",
]: add_bullet(doc, item)

heading(doc, "4. Proposed End-to-End Workflow")
add_note(doc, "Validation label", "The sequence below is a proposed operating model derived from the planning documents. Jeff must confirm the actual stages, entry criteria, exit criteria, and HubSpot mapping.")
workflow = [
    ("1", "Prospect identified", "Lead source and ownership recorded"),
    ("2", "Company qualified", "ICP rules and disqualifiers applied"),
    ("3", "Decision maker identified", "Primary and backup contacts validated"),
    ("4", "Research completed", "Evidence, confidence, and gaps documented"),
    ("5", "Outreach prepared", "Channel, message, CTA, and follow-up plan drafted"),
    ("6", "Human approval", "Jeff reviews before any external communication"),
    ("7", "Outreach sent and logged", "Activity, date, source, and next step captured"),
    ("8", "Response classified", "Affirmative, objection, timing, referral, or stop"),
    ("9", "Conversation / assessment", "Briefing, notes, actions, and follow-up drafted"),
    ("10", "Opportunity through client", "Pipeline, attribution, outcome, and commission evidence maintained"),
]
add_table(doc, ["Stage", "Working Stage", "Required Output"], workflow, [840, 3000, 5520])

heading(doc, "5. Proposed Nine-Agent Operating Model")
add_note(doc, "Design principle", "Treat these as modular roles first. Do not build nine separate complex automations until the shared workflow, data contract, and approval rules are validated.")
agents = [
    ("1", "Prospect Research", "Company fit, evidence, disqualifiers, source links", "Jeff validates rules and exceptions"),
    ("2", "Decision Maker", "Primary/secondary contacts and confidence", "Jeff resolves ambiguous titles or employment"),
    ("3", "Relationship Mapping", "Warm paths and appropriate connection angles", "Jeff approves introduction requests"),
    ("4", "Company Intelligence", "Concise account brief and possible spend categories", "No unsupported overspending claims"),
    ("5", "Outreach", "Email, LinkedIn, call, voicemail, and follow-up drafts", "Jeff approves every send in Phase 1"),
    ("6", "CRM Management", "Prepared records, notes, tasks, and attribution updates", "Sensitive writes require approval"),
    ("7", "Meeting Preparation", "Account briefing, questions, objections, talking points", "Jeff selects final agenda"),
    ("8", "Post-Meeting", "Summary, follow-up draft, tasks, next agenda", "Jeff verifies notes and approves follow-up"),
    ("9", "Pipeline Coach", "Daily priorities, stalled deals, next actions, weekly review", "Jeff owns prioritization and disposition"),
]
add_table(doc, ["#", "Role", "Core Output", "Human Checkpoint"], agents, [420, 1720, 3740, 3480])

heading(doc, "6. Proposed Scoring and Prioritization Model")
doc.add_paragraph("Use this only as a discussion starter. Jeff must define the factors, weights, thresholds, and examples before the score drives any action.")
score_rows = [
    ("Company fit", "25", "Industry, geography, ownership, locations, operating model"),
    ("Company size", "20", "Revenue, employees, footprint, estimated addressable spend"),
    ("Potential cost opportunity", "20", "Relevant expense categories and observable business signals"),
    ("Decision-maker access", "15", "Contact confidence, reachable finance/operations leader"),
    ("Relationship strength", "10", "Direct relationship, warm introduction, credible mutual connection"),
    ("Current business signal", "10", "Expansion, acquisition, leadership change, hiring, or other relevant trigger"),
]
add_table(doc, ["Factor", "Draft Points", "Evidence to Define"], score_rows, [2460, 1260, 5640])
add_bullet(doc, "Proposed pursue threshold: Jeff to define.")
add_bullet(doc, "Automatic disqualifiers: Jeff to define.")
add_bullet(doc, "Low-confidence rule: route to Jeff rather than infer missing facts.")

heading(doc, "7. Non-Negotiable Guardrails")
for item in [
    "Never claim or imply that a company is overspending without verified evidence.",
    "Never reveal intent-data sources or other sensitive research methods in outreach.",
    "Never fabricate a business problem, mutual relationship, testimonial, result, or personal detail.",
    "Never use personal information that would feel invasive, irrelevant, or inappropriate.",
    "Never send, schedule, publish, modify sensitive CRM fields, or contact an existing client without the approved human checkpoint.",
    "When sources disagree, preserve the conflict, show confidence, and ask Jeff to decide.",
    "Log what the system recommended or changed, the source evidence, and who approved it.",
    "Honor opt-outs, do-not-contact requirements, channel rules, and applicable outreach/recording policies.",
]: add_bullet(doc, item)

heading(doc, "8. Jeff’s Follow-Up Submission Package")
doc.add_paragraph("Jeff can submit the material in batches. Original files are better than summaries because the examples will be used to extract voice, rules, edge cases, and CRM requirements.")
submission = [
    ("A — Required first", "3–5 real target companies; 3 past wins; 3 strong-looking losses; ICP examples and disqualifiers", "Unlocks qualification and scoring rules"),
    ("B — Messaging evidence", "5 successful prospecting emails; 5 weak/unsuccessful emails; typical follow-ups; LinkedIn conversations; scripts", "Unlocks Jeff voice and outreach playbook"),
    ("C — Offer evidence", "Spend Assessment materials; service descriptions; presentation; FAQs; objections; case studies; testimonials", "Unlocks accurate offer messaging"),
    ("D — CRM evidence", "HubSpot screenshots/export of pipeline stages, lead statuses, required fields, sequences, tasks, and current automations", "Unlocks data mapping and safe write rules"),
    ("E — Commercial rules", "Attribution definition; commission qualification; existing-relationship treatment; reviewer and review cadence", "Unlocks auditable revenue-share tracking"),
    ("F — Approval rules", "What AI may draft, create, update, schedule, or send; what always requires approval; unacceptable-error list", "Unlocks autonomy policy"),
]
add_table(doc, ["Batch", "Items", "What It Unlocks"], submission, [1680, 4860, 2820])

heading(doc, "9. Live Workflow Capture Instructions")
doc.add_paragraph("Record a screen-share while Jeff completes the real process from start to finish using HubSpot, Sales Navigator, ZoomInfo, LeadIQ, company sites, email, and any other actual tools. Jeff should narrate why he clicks, skips, trusts, rejects, or changes course.")
for item in [
    "One prospect Jeff accepts quickly.",
    "One prospect Jeff rejects after research.",
    "One ambiguous or surprising prospect that exposes an exception.",
    "One outreach draft and the edits Jeff makes before approving it.",
    "One HubSpot update, including the fields he always completes and the fields he avoids.",
    "One follow-up decision: continue, change channel, nurture, refer, or stop.",
]: add_number(doc, item)
heading(doc, "For every meaningful decision, capture", 2)
add_table(doc, ["Field", "Question"], [
    ("Trigger", "What caused Jeff to consider taking action?"),
    ("Decision", "What did Jeff decide?"),
    ("Rule", "Why did he make that decision, and how could the system repeat it?"),
    ("Action", "What happened next, in which tool, and what was recorded?"),
    ("Exception", "When would Jeff do something different?"),
    ("Confidence", "What evidence is sufficient, and when must the system ask?"),
], [1800, 7560])

heading(doc, "10. Open Decision Register")
open_items = [
    ("ICP", "Minimum/ideal size, industries, geography, ownership, locations, disqualifiers", "Jeff"),
    ("Offer", "Exact Spend Assessment promise, inputs, outputs, timing, obligation, next steps", "Jeff"),
    ("Scoring", "Factors, weights, pursue threshold, stop threshold, confidence policy", "Jeff + Coach Bouf"),
    ("Contacts", "Preferred titles by company size and source-trust hierarchy", "Jeff"),
    ("Messaging", "Voice, CTA, length, cadence, stop rules, phrases to use/avoid", "Jeff"),
    ("HubSpot", "Pipeline stages, required properties, permissions, sequences, integrations", "Jeff + system review"),
    ("Attribution", "Sourced vs. assisted, affirmative-interest signal, existing relationships, verification", "Jeff + Coach Bouf"),
    ("Autonomy", "Draft/create/update/send/schedule permissions and approvals", "Jeff + Coach Bouf"),
    ("Compliance", "Outreach channels, consent, opt-out, recording, retention, and data-use rules", "Business/legal owner"),
    ("Success", "30-day targets and baseline measures", "Jeff + Coach Bouf"),
]
add_table(doc, ["Area", "Decision Needed", "Owner"], open_items, [1680, 6120, 1560])

heading(doc, "11. Recommended Build Sequence")
for item in [
    "Add the verified Interview 1 transcript and extract decisions, examples, contradictions, and unresolved questions.",
    "Receive Jeff’s first follow-up batch: real prospects, wins/losses, and messaging examples.",
    "Run the live screen-share workflow capture and validate the actual tool sequence.",
    "Finalize the sales playbook, ICP rules, scoring rubric, guardrails, and HubSpot data contract.",
    "Build one approval-gated pilot workflow: research → recommendation → outreach draft → Jeff approval → CRM logging.",
    "Test on three to five prospects and compare manual versus AI-assisted time, accuracy, and quality.",
    "Only then decide which modular roles deserve separate skills or automations.",
]: add_number(doc, item)

heading(doc, "12. Phase 1 Acceptance Criteria")
for item in [
    "Jeff confirms that the documented workflow matches how he actually works.",
    "Every qualification and scoring rule has examples and a documented exception path.",
    "Every generated factual claim is traceable to a source or labeled as a hypothesis.",
    "No external communication occurs without the approved checkpoint.",
    "HubSpot field mapping and attribution definitions are approved before automation writes.",
    "Three to five real-prospect tests are completed and reviewed.",
    "Baseline and pilot measures show whether the workflow saves time without lowering quality or trust.",
]: add_bullet(doc, item)

heading(doc, "Appendix A — Recommended HubSpot Attribution Fields")
add_note(doc, "Status", "These fields came from the session plan and require validation against the current HubSpot schema before creation.")
add_table(doc, ["Field", "Working Purpose"], [
    ("Outreach Engine Sourced", "Yes / No / Assisted / Unknown classification"),
    ("Outreach Engine Assisted", "Records material AI support on an existing or human-sourced opportunity"),
    ("First AI Outreach Date", "Date of the first approved AI-assisted outreach"),
    ("First Affirmative Interest Date", "Date of the agreed qualifying response or action"),
    ("Affirmative Interest Source", "Email, LinkedIn, phone, meeting, referral, or other approved source"),
    ("Original Lead Source", "Preserves the original acquisition source"),
    ("AI Agent Used", "Records the role or workflow that contributed"),
    ("Qualified for Revenue Share", "Approval-controlled commercial determination"),
], [3420, 5940])

heading(doc, "Appendix B — Consolidation Notes")
doc.add_paragraph("The questionnaires overlap: retain the longer version as the master bank and the shorter version as the interviewer guide. The 120-minute plan proposes structure, but it is not proof of Jeff’s actual process. This brief preserves all unknowns as open decisions.")

doc.core_properties.title = "Jeff Outreach Engine Consolidated Discovery Brief"
doc.core_properties.subject = "Discovery consolidation and follow-up submission guide"
doc.core_properties.author = "AI Growth Plan"
doc.core_properties.keywords = "Jeff Outreach Engine, discovery, HubSpot, prospecting, AI agents, approval"
doc.save(OUT)
print(OUT.resolve())
