from pathlib import Path

from docx import Document


for path in sorted(Path("docs").rglob("*.docx")):
    print(f"\n### FILE: {path.name}")
    document = Document(path)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            print(text)
    for table_index, table in enumerate(document.tables, start=1):
        print(f"TABLE {table_index}")
        for row in table.rows:
            print(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))
